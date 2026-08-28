# -*- coding: utf-8 -*-
"""Exportadores de un conjunto de experiencias seleccionadas a Word y Excel.

Pensado para el flujo de postulación a fondos/proyectos: el usuario filtra en el
Explorador, marca las experiencias relevantes y se lleva una ficha en Word (para pegar
en una propuesta) y una planilla Excel (para anexos / trabajo de datos).
"""
import io
from datetime import date

import pandas as pd

# (etiqueta visible, columna en el DataFrame de noticias). Orden pensado para una propuesta.
FICHA_FIELDS = [
    ("Resumen", "descripcion_catalogo"),
    ("Categoría macro", "categoria_macro"),
    ("Categoría temática", "categorias"),
    ("Metodología de facilitación", "metodologia"),
    ("Actores institucionales", "actores_normalizados"),
    ("Eje GCAA (acción climática)", "eje_gcaa"),
    ("Objetivo GCAA", "objetivo_gcaa"),
    ("Atributo de resiliencia (CR2)", "atributos_resiliencia"),
    ("Sub-atributo de resiliencia", "subatributos_resiliencia"),
    ("Beneficiarios directos", "beneficiarios_directos"),
    ("Beneficiarios indirectos", "beneficiarios_indirectos"),
    ("Enfoque de género", "enfoque_genero"),
]

# Columnas para la planilla Excel "Resumen" (clave para anexos de postulación).
EXCEL_KEY_COLS = [
    ("titulo", "Título"),
    ("anio", "Año"),
    ("pais", "País"),
    ("lugar", "Lugar"),
    ("descripcion_catalogo", "Resumen"),
    ("categoria_macro", "Categoría macro"),
    ("categorias", "Categoría temática"),
    ("metodologia", "Metodología"),
    ("actores_normalizados", "Actores institucionales"),
    ("eje_gcaa", "Eje GCAA"),
    ("objetivo_gcaa", "Objetivo GCAA"),
    ("atributos_resiliencia", "Atributo de resiliencia"),
    ("subatributos_resiliencia", "Sub-atributo de resiliencia"),
    ("beneficiarios_directos", "Beneficiarios directos"),
    ("beneficiarios_indirectos", "Beneficiarios indirectos"),
    ("enfoque_genero", "Enfoque de género"),
    ("url_noticia", "Enlace"),
    ("contenido_completo", "Texto completo"),
]

_EMPTY = {"", "no aplica", "no especificado", "sin dato", "nan", "none"}


def _txt(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value).strip()


def _clean_multi(value) -> str:
    """'a;b ; c' -> 'a; b; c', quitando vacíos y marcadores de 'sin valor'."""
    raw = _txt(value)
    if not raw:
        return ""
    parts = [p.strip() for p in raw.split(";")]
    parts = [p for p in parts if p and p.lower() not in _EMPTY]
    return "; ".join(dict.fromkeys(parts))


def _anio(row) -> str:
    for key in ("anio", "fecha_parsed", "fecha_publicacion"):
        v = row.get(key)
        if v is None or (isinstance(v, float) and pd.isna(v)):
            continue
        if key == "anio":
            try:
                return str(int(v))
            except (TypeError, ValueError):
                continue
        if key == "fecha_parsed" and not pd.isna(v):
            return str(pd.Timestamp(v).year)
        return _txt(v)
    return "s/f"


# --------------------------------------------------------------------------- Excel
def experiences_to_excel(rows: pd.DataFrame) -> bytes:
    """Planilla con dos hojas: 'Resumen' (columnas clave) y 'Datos completos' (todo)."""
    resumen = pd.DataFrame({
        label: rows[col].map(_clean_multi) if col in rows.columns else ""
        for col, label in EXCEL_KEY_COLS
    })

    completo = rows.drop(
        columns=[c for c in rows.columns if c.endswith("_primary")]
        + [c for c in ("item", "enlaces_externos_lista", "fecha_parsed", "tiene_fecha") if c in rows.columns],
        errors="ignore",
    ).copy()
    for c in completo.columns:
        completo[c] = completo[c].map(_txt)

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as xls:
        resumen.to_excel(xls, sheet_name="Resumen", index=False)
        completo.to_excel(xls, sheet_name="Datos completos", index=False)
        for sheet in xls.book.worksheets:
            for column_cells in sheet.columns:
                width = min(60, max(12, max((len(str(c.value or "")) for c in column_cells[:40]), default=12) + 2))
                sheet.column_dimensions[column_cells[0].column_letter].width = width
    return buf.getvalue()


# ---------------------------------------------------------------------------- Word
def experiences_to_word(rows: pd.DataFrame, contexto: str = "") -> bytes:
    """Documento con una portada + una ficha por experiencia (info clave para postular)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("Experiencias seleccionadas — Catálogo Glocalminds", level=0)
    p = doc.add_paragraph()
    p.add_run(f"{len(rows)} experiencia(s)  ·  Generado el {date.today().strftime('%d-%m-%Y')}").italic = True
    if contexto.strip():
        doc.add_paragraph(f"Contexto de la selección: {contexto.strip()}")

    # Índice / tabla resumen
    doc.add_heading("Resumen", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, head in enumerate(("Título", "Año", "País", "Categoría macro")):
        table.rows[0].cells[i].text = head
    for _, row in rows.iterrows():
        cells = table.add_row().cells
        cells[0].text = _txt(row.get("titulo"))
        cells[1].text = _anio(row)
        cells[2].text = _clean_multi(row.get("pais")) or "s/d"
        cells[3].text = _clean_multi(row.get("categoria_macro")) or "s/d"

    # Una ficha por experiencia
    for _, row in rows.iterrows():
        doc.add_page_break()
        doc.add_heading(_txt(row.get("titulo")) or "Sin título", level=1)

        meta = "  ·  ".join(
            x for x in (
                f"Año: {_anio(row)}",
                f"País: {_clean_multi(row.get('pais'))}" if _clean_multi(row.get("pais")) else "",
                f"Lugar: {_clean_multi(row.get('lugar'))}" if _clean_multi(row.get("lugar")) else "",
            ) if x
        )
        doc.add_paragraph(meta).runs[0].italic = True

        for label, col in FICHA_FIELDS:
            val = _clean_multi(row.get(col))
            if not val:
                continue
            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            run.bold = True
            para.add_run(val)

        url = _txt(row.get("url_noticia"))
        if url:
            para = doc.add_paragraph()
            para.add_run("Enlace: ").bold = True
            para.add_run(url)

        texto = _txt(row.get("contenido_completo"))
        if texto and texto.lower() != _txt(row.get("descripcion_catalogo")).lower():
            doc.add_heading("Texto completo", level=2)
            for bloque in texto.split("\n"):
                if bloque.strip():
                    doc.add_paragraph(bloque.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
